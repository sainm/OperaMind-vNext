"""Bounded XLSX and DOCX structural-signal extraction."""

from __future__ import annotations

import re
import zipfile
from collections.abc import Collection, Iterable, Mapping
from dataclasses import dataclass
from datetime import date, datetime, time, timedelta
from importlib.metadata import version as distribution_version
from itertools import pairwise
from math import isfinite
from pathlib import Path, PurePosixPath
from typing import Protocol, runtime_checkable
from xml.etree.ElementTree import ParseError

from docx import Document
from docx.opc.exceptions import OpcError
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.utils.exceptions import InvalidFileException

from operamind.domain.canonical_facts import (
    ObservedField,
    ObservedRecord,
    normalize_business_value,
    normalize_field_name,
)
from operamind.domain.document_conventions import ConventionVariant, DocumentSignals

HTTP_OPERATION = re.compile(
    r"^(?:GET|POST|PUT|PATCH|DELETE|HEAD|OPTIONS)\s+/\S+$",
    re.IGNORECASE,
)


class OfficeDocumentError(ValueError):
    """Base error for rejected or unreadable Office source documents."""


class OfficeDocumentSecurityError(OfficeDocumentError):
    """Raised before parsing when an Office ZIP violates configured limits."""


class UnsupportedDocumentTypeError(OfficeDocumentError):
    """Raised when no explicitly registered extractor supports a suffix."""


@dataclass(frozen=True, slots=True)
class ExtractionLimits:
    """Resource and scan bounds applied before and during Office parsing."""

    max_file_bytes: int = 50 * 1024 * 1024
    max_archive_entries: int = 10_000
    max_uncompressed_bytes: int = 200 * 1024 * 1024
    max_scan_rows: int = 500
    max_scan_columns: int = 100
    max_paragraphs: int = 2_000

    def __post_init__(self) -> None:
        field_names = (
            "max_file_bytes",
            "max_archive_entries",
            "max_uncompressed_bytes",
            "max_scan_rows",
            "max_scan_columns",
            "max_paragraphs",
        )
        for field_name in field_names:
            if getattr(self, field_name) <= 0:
                raise ValueError(f"{field_name} must be greater than zero")


@dataclass(frozen=True, slots=True)
class _DetectedHeaders:
    by_column: Mapping[int, str]
    canonical_fields: frozenset[str]


class DocumentSignalExtractor(Protocol):
    """Port implemented by each supported document-format adapter."""

    @property
    def suffixes(self) -> frozenset[str]: ...

    @property
    def extractor_ref(self) -> str: ...

    def extract(self, path: Path) -> DocumentSignals: ...

    def extract_records(
        self, path: Path, variant: ConventionVariant
    ) -> tuple[ObservedRecord, ...]: ...


@runtime_checkable
class WorksheetDocumentSignalExtractor(DocumentSignalExtractor, Protocol):
    """Optional capabilities exposed by worksheet-based document formats."""

    def extract_sheet_signals(self, path: Path) -> tuple[tuple[str, DocumentSignals], ...]: ...

    def extract_records_for_sheet(
        self,
        path: Path,
        variant: ConventionVariant,
        *,
        sheet_name: str,
    ) -> tuple[ObservedRecord, ...]: ...


@runtime_checkable
class LearningStructureSignalExtractor(DocumentSignalExtractor, Protocol):
    """Optional structure-only view that excludes changing business row values."""

    def extract_learning_structure(
        self, path: Path
    ) -> tuple[tuple[str, DocumentSignals], ...]: ...


class DocumentSignalExtractorRegistry:
    """Select a format adapter by exact, case-insensitive file suffix."""

    def __init__(self, extractors: Iterable[DocumentSignalExtractor]) -> None:
        by_suffix: dict[str, DocumentSignalExtractor] = {}
        for extractor in extractors:
            for suffix in extractor.suffixes:
                normalized = suffix.casefold()
                if not normalized.startswith("."):
                    raise ValueError(f"Extractor suffix must start with '.': {suffix}")
                if normalized in by_suffix:
                    raise ValueError(f"Duplicate document extractor suffix: {normalized}")
                by_suffix[normalized] = extractor
        self._by_suffix = by_suffix

    @classmethod
    def default(cls, limits: ExtractionLimits | None = None) -> DocumentSignalExtractorRegistry:
        """Build the MVP registry for modern, ZIP-based Office formats."""

        effective_limits = limits or ExtractionLimits()
        return cls(
            (
                XlsxSignalExtractor(effective_limits),
                DocxSignalExtractor(effective_limits),
            )
        )

    def extract(self, path: Path) -> DocumentSignals:
        """Extract signals or reject unsupported legacy/binary formats explicitly."""

        return self._extractor_for(path).extract(path)

    def extract_records(self, path: Path, variant: ConventionVariant) -> tuple[ObservedRecord, ...]:
        """Extract alias-bearing table rows for one already selected Variant."""

        extractor = self._extractor_for(path)
        return extractor.extract_records(path, variant)

    def extract_sheet_signals(self, path: Path) -> tuple[tuple[str, DocumentSignals], ...]:
        """Return bounded per-sheet signals when the format exposes worksheets."""

        extractor = self._extractor_for(path)
        if not isinstance(extractor, WorksheetDocumentSignalExtractor):
            return ()
        return extractor.extract_sheet_signals(path)

    def extract_learning_structure(
        self, path: Path
    ) -> tuple[tuple[str, DocumentSignals], ...]:
        """Return headings and header rows without mutable business data."""

        extractor = self._extractor_for(path)
        if not isinstance(extractor, LearningStructureSignalExtractor):
            return ((path.name, extractor.extract(path)),)
        return extractor.extract_learning_structure(path)

    def extract_records_for_sheet(
        self,
        path: Path,
        variant: ConventionVariant,
        *,
        sheet_name: str,
    ) -> tuple[ObservedRecord, ...]:
        """Extract one Variant only from one named worksheet."""

        extractor = self._extractor_for(path)
        if not isinstance(extractor, WorksheetDocumentSignalExtractor):
            return extractor.extract_records(path, variant)
        return extractor.extract_records_for_sheet(path, variant, sheet_name=sheet_name)

    def extractor_ref(self, path: Path) -> str:
        """Return the implementation and parser-library version selected for one path."""

        return self._extractor_for(path).extractor_ref

    @property
    def supported_suffixes(self) -> frozenset[str]:
        """Return the exact, normalized file suffixes accepted by this registry."""

        return frozenset(self._by_suffix)

    def _extractor_for(self, path: Path) -> DocumentSignalExtractor:
        extractor = self._by_suffix.get(path.suffix.casefold())
        if extractor is None:
            supported = ", ".join(sorted(self._by_suffix))
            raise UnsupportedDocumentTypeError(
                f"Unsupported document type {path.suffix or '<none>'}; supported: {supported}"
            )
        return extractor


@dataclass(frozen=True, slots=True)
class XlsxSignalExtractor:
    """Extract sheet names and bounded textual table structure from XLSX."""

    limits: ExtractionLimits

    @property
    def suffixes(self) -> frozenset[str]:
        return frozenset({".xlsx"})

    @property
    def extractor_ref(self) -> str:
        return f"operamind-xlsx-structural@1+openpyxl-{distribution_version('openpyxl')}"

    def extract(self, path: Path) -> DocumentSignals:
        checked_path = _validate_office_archive(path, self.limits)
        try:
            workbook = load_workbook(
                checked_path,
                read_only=True,
                data_only=False,
                keep_links=False,
            )
        except (
            OSError,
            ValueError,
            KeyError,
            SyntaxError,
            ParseError,
            InvalidFileException,
            zipfile.BadZipFile,
        ) as error:
            raise OfficeDocumentError(f"Cannot parse XLSX document: {checked_path}") from error

        headings: list[str] = []
        headers: list[str] = []
        business_terms: list[str] = []
        try:
            sheet_names = tuple(workbook.sheetnames)
            for worksheet in workbook.worksheets:
                rows = worksheet.iter_rows(
                    min_row=1,
                    max_row=self.limits.max_scan_rows,
                    max_col=self.limits.max_scan_columns,
                    values_only=True,
                )
                for row in rows:
                    values = _text_values(row)
                    business_terms.extend(values)
                    if len(values) == 1:
                        headings.extend(values)
                    elif len(values) >= 2:
                        headers.extend(values)
        finally:
            workbook.close()

        return DocumentSignals.from_raw(
            filename=checked_path.name,
            sheet_names=sheet_names,
            headings=tuple(headings),
            headers=tuple(headers),
            business_terms=tuple(business_terms),
        )

    def extract_sheet_signals(self, path: Path) -> tuple[tuple[str, DocumentSignals], ...]:
        """Extract independent structural signals for every bounded worksheet."""

        checked_path = _validate_office_archive(path, self.limits)
        try:
            workbook = load_workbook(
                checked_path,
                read_only=True,
                data_only=False,
                keep_links=False,
            )
        except (
            OSError,
            ValueError,
            KeyError,
            SyntaxError,
            ParseError,
            InvalidFileException,
            zipfile.BadZipFile,
        ) as error:
            raise OfficeDocumentError(f"Cannot parse XLSX document: {checked_path}") from error

        result: list[tuple[str, DocumentSignals]] = []
        try:
            for worksheet in workbook.worksheets:
                headings: list[str] = []
                headers: list[str] = []
                business_terms: list[str] = []
                for row in worksheet.iter_rows(
                    min_row=1,
                    max_row=self.limits.max_scan_rows,
                    max_col=self.limits.max_scan_columns,
                    values_only=True,
                ):
                    values = _text_values(row)
                    business_terms.extend(values)
                    if len(values) == 1:
                        headings.extend(values)
                    elif len(values) >= 2:
                        headers.extend(values)
                result.append(
                    (
                        worksheet.title,
                        DocumentSignals.from_raw(
                            filename=checked_path.name,
                            sheet_names=(worksheet.title,),
                            headings=tuple(headings),
                            headers=tuple(headers),
                            business_terms=tuple(business_terms),
                        ),
                    )
                )
        finally:
            workbook.close()
        return tuple(result)

    def extract_learning_structure(
        self, path: Path
    ) -> tuple[tuple[str, DocumentSignals], ...]:
        checked_path = _validate_office_archive(path, self.limits)
        try:
            workbook = load_workbook(
                checked_path,
                read_only=True,
                data_only=False,
                keep_links=False,
            )
        except (
            OSError,
            ValueError,
            KeyError,
            SyntaxError,
            ParseError,
            InvalidFileException,
            zipfile.BadZipFile,
        ) as error:
            raise OfficeDocumentError(f"Cannot parse XLSX document: {checked_path}") from error
        result: list[tuple[str, DocumentSignals]] = []
        try:
            for worksheet in workbook.worksheets:
                headings: list[str] = []
                headers: tuple[str, ...] = ()
                for row in worksheet.iter_rows(
                    min_row=1,
                    max_row=self.limits.max_scan_rows,
                    max_col=self.limits.max_scan_columns,
                    values_only=True,
                ):
                    values = _text_values(row)
                    if not values:
                        continue
                    if len(values) == 1 and not headers:
                        headings.extend(values)
                        continue
                    if len(values) >= 2:
                        headers = tuple(values)
                        break
                result.append(
                    (
                        worksheet.title,
                        DocumentSignals.from_raw(
                            filename=checked_path.name,
                            sheet_names=(worksheet.title,),
                            headings=tuple(headings),
                            headers=headers,
                        ),
                    )
                )
        finally:
            workbook.close()
        return tuple(result)

    def extract_records(self, path: Path, variant: ConventionVariant) -> tuple[ObservedRecord, ...]:
        """Find Variant header rows and retain row/cell source locations."""

        checked_path = _validate_office_archive(path, self.limits)
        try:
            workbook = load_workbook(
                checked_path,
                read_only=True,
                data_only=False,
                keep_links=False,
            )
        except (
            OSError,
            ValueError,
            KeyError,
            SyntaxError,
            ParseError,
            InvalidFileException,
            zipfile.BadZipFile,
        ) as error:
            raise OfficeDocumentError(f"Cannot parse XLSX document: {checked_path}") from error

        records: list[ObservedRecord] = []
        try:
            context_rows = (
                (
                    tuple(cell.value for cell in row),
                    tuple(
                        f"{checked_path.name}#{worksheet.title}!"
                        f"{get_column_letter(column_index)}{row_index}"
                        for column_index, _cell in enumerate(row, start=1)
                    ),
                )
                for worksheet in workbook.worksheets
                for row_index, row in enumerate(
                    worksheet.iter_rows(
                        min_row=1,
                        max_row=self.limits.max_scan_rows,
                        max_col=self.limits.max_scan_columns,
                        values_only=False,
                    ),
                    start=1,
                )
            )
            context_fields = _key_value_context_fields(context_rows, variant)
            for worksheet in workbook.worksheets:
                active_headers: _DetectedHeaders | None = None
                rows = worksheet.iter_rows(
                    min_row=1,
                    max_row=self.limits.max_scan_rows,
                    max_col=self.limits.max_scan_columns,
                    values_only=False,
                )
                for row_index, row in enumerate(rows, start=1):
                    raw_values = tuple(_text_value(cell.value) for cell in row)
                    detected_headers = _recognized_headers(
                        raw_values,
                        variant,
                        context_fields.keys(),
                    )
                    if detected_headers is not None:
                        active_headers = detected_headers
                        continue
                    if active_headers is None:
                        continue
                    fields = tuple(
                        ObservedField(
                            name=header,
                            value=value,
                            source_ref=(
                                f"{checked_path.name}#{worksheet.title}!{row[column_index].coordinate}"
                            ),
                        )
                        for column_index, header in active_headers.by_column.items()
                        if (value := _record_value(row[column_index].value)) is not None
                    )
                    inherited_fields = tuple(
                        field
                        for canonical_field, fields_for_canonical in context_fields.items()
                        if canonical_field not in active_headers.canonical_fields
                        for field in fields_for_canonical
                    )
                    fields = inherited_fields + fields
                    if fields:
                        records.append(
                            ObservedRecord(
                                record_ref=(
                                    f"{checked_path.name}#{worksheet.title}!row={row_index}"
                                ),
                                fields=fields,
                            )
                        )
        finally:
            workbook.close()
        return tuple(records)

    def extract_records_for_sheet(
        self,
        path: Path,
        variant: ConventionVariant,
        *,
        sheet_name: str,
    ) -> tuple[ObservedRecord, ...]:
        """Extract records for one worksheet without mixing table conventions."""

        checked_path = _validate_office_archive(path, self.limits)
        if not sheet_name.strip():
            raise ValueError("sheet_name must not be blank")
        try:
            workbook = load_workbook(
                checked_path,
                read_only=True,
                data_only=False,
                keep_links=False,
            )
        except (
            OSError,
            ValueError,
            KeyError,
            SyntaxError,
            ParseError,
            InvalidFileException,
            zipfile.BadZipFile,
        ) as error:
            raise OfficeDocumentError(f"Cannot parse XLSX document: {checked_path}") from error

        try:
            worksheet = workbook[sheet_name]
        except KeyError as error:
            workbook.close()
            raise OfficeDocumentError(
                f"Worksheet does not exist in XLSX document: {sheet_name}"
            ) from error

        try:
            all_context_rows = tuple(
                (
                    tuple(cell.value for cell in row),
                    tuple(
                        f"{checked_path.name}#{context_sheet.title}!"
                        f"{get_column_letter(column_index)}{row_index}"
                        for column_index, _cell in enumerate(row, start=1)
                    ),
                )
                for context_sheet in workbook.worksheets
                for row_index, row in enumerate(
                    context_sheet.iter_rows(
                        min_row=1,
                        max_row=self.limits.max_scan_rows,
                        max_col=self.limits.max_scan_columns,
                        values_only=False,
                    ),
                    start=1,
                )
            )
            context_fields = _key_value_context_fields(all_context_rows, variant)
            local_rows = tuple(
                (
                    tuple(cell.value for cell in row),
                    tuple(
                        f"{checked_path.name}#{worksheet.title}!"
                        f"{get_column_letter(column_index)}{row_index}"
                        for column_index, _cell in enumerate(row, start=1)
                    ),
                )
                for row_index, row in enumerate(
                    worksheet.iter_rows(
                        min_row=1,
                        max_row=self.limits.max_scan_rows,
                        max_col=self.limits.max_scan_columns,
                        values_only=False,
                    ),
                    start=1,
                )
            )
            local_key_value_fields = _key_value_fields(local_rows, variant)
            effective_context_fields = dict(context_fields)
            effective_context_fields.update(
                {
                    canonical_field: fields
                    for canonical_field, fields in local_key_value_fields.items()
                    if canonical_field in variant.stable_key_fields
                }
            )
            active_headers: _DetectedHeaders | None = None
            active_section_fields: dict[str, ObservedField] = {}
            records: list[ObservedRecord] = []
            rows = worksheet.iter_rows(
                min_row=1,
                max_row=self.limits.max_scan_rows,
                max_col=self.limits.max_scan_columns,
                values_only=False,
            )
            for row_index, row in enumerate(rows, start=1):
                raw_values = tuple(_text_value(cell.value) for cell in row)
                section_fields = _section_marker_fields(
                    raw_values,
                    tuple(
                        f"{checked_path.name}#{worksheet.title}!"
                        f"{get_column_letter(column_index)}{row_index}"
                        for column_index, _cell in enumerate(row, start=1)
                    ),
                    variant,
                    excluded_fields=frozenset(effective_context_fields),
                )
                if section_fields:
                    active_section_fields.update(section_fields)
                    active_headers = None
                    continue
                detected_headers = _recognized_headers(
                    raw_values,
                    variant,
                    effective_context_fields.keys() | active_section_fields.keys(),
                )
                if detected_headers is not None:
                    active_headers = detected_headers
                    continue
                if active_headers is None:
                    continue
                fields = tuple(
                    ObservedField(
                        name=header,
                        value=value,
                        source_ref=(
                            f"{checked_path.name}#{worksheet.title}!{row[column_index].coordinate}"
                        ),
                    )
                    for column_index, header in active_headers.by_column.items()
                    if (value := _record_value(row[column_index].value)) is not None
                )
                inherited_fields = tuple(
                    field
                    for canonical_field, fields_for_canonical in effective_context_fields.items()
                    if canonical_field not in active_headers.canonical_fields
                    for field in fields_for_canonical
                ) + tuple(
                    field
                    for canonical_field, field in active_section_fields.items()
                    if canonical_field not in active_headers.canonical_fields
                )
                fields = inherited_fields + fields
                if fields:
                    records.append(
                        ObservedRecord(
                            record_ref=f"{checked_path.name}#{worksheet.title}!row={row_index}",
                            fields=fields,
                        )
                    )
            if not records and set(variant.stable_key_fields).issubset(local_key_value_fields):
                records.append(
                    ObservedRecord(
                        record_ref=f"{checked_path.name}#{worksheet.title}!key-values",
                        fields=tuple(
                            field for fields in local_key_value_fields.values() for field in fields
                        ),
                    )
                )
            return tuple(records)
        finally:
            workbook.close()


@dataclass(frozen=True, slots=True)
class DocxSignalExtractor:
    """Extract headings, table headers, and terms from DOCX."""

    limits: ExtractionLimits

    @property
    def suffixes(self) -> frozenset[str]:
        return frozenset({".docx"})

    @property
    def extractor_ref(self) -> str:
        return f"operamind-docx-structural@1+python-docx-{distribution_version('python-docx')}"

    def extract(self, path: Path) -> DocumentSignals:
        checked_path = _validate_office_archive(path, self.limits)
        try:
            document = Document(str(checked_path))
        except (
            OSError,
            ValueError,
            KeyError,
            SyntaxError,
            ParseError,
            OpcError,
            zipfile.BadZipFile,
        ) as error:
            raise OfficeDocumentError(f"Cannot parse DOCX document: {checked_path}") from error

        headings: list[str] = []
        headers: list[str] = []
        business_terms: list[str] = []
        for paragraph in document.paragraphs[: self.limits.max_paragraphs]:
            value = _text_value(paragraph.text)
            if value is None:
                continue
            business_terms.append(value)
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.casefold().startswith("heading") or "見出し" in style_name:
                headings.append(value)

        for table in document.tables:
            for row_index, row in enumerate(table.rows[: self.limits.max_scan_rows]):
                values = _text_values(
                    cell.text for cell in row.cells[: self.limits.max_scan_columns]
                )
                business_terms.extend(values)
                if row_index == 0:
                    headers.extend(values)

        return DocumentSignals.from_raw(
            filename=checked_path.name,
            headings=tuple(headings),
            headers=tuple(headers),
            business_terms=tuple(business_terms),
        )

    def extract_records(self, path: Path, variant: ConventionVariant) -> tuple[ObservedRecord, ...]:
        """Extract rows from DOCX tables whose first row identifies Stable Key fields."""

        checked_path = _validate_office_archive(path, self.limits)
        try:
            document = Document(str(checked_path))
        except (
            OSError,
            ValueError,
            KeyError,
            SyntaxError,
            ParseError,
            OpcError,
            zipfile.BadZipFile,
        ) as error:
            raise OfficeDocumentError(f"Cannot parse DOCX document: {checked_path}") from error

        records: list[ObservedRecord] = []
        context_rows = (
            (
                tuple(cell.text for cell in row.cells[: self.limits.max_scan_columns]),
                tuple(
                    (
                        f"{checked_path.name}#table={table_index},row={row_index},"
                        f"column={column_index}"
                    )
                    for column_index, _cell in enumerate(
                        row.cells[: self.limits.max_scan_columns], start=1
                    )
                ),
            )
            for table_index, table in enumerate(document.tables, start=1)
            for row_index, row in enumerate(table.rows[: self.limits.max_scan_rows], start=1)
        )
        context_fields = _key_value_context_fields(context_rows, variant)
        for table_index, table in enumerate(document.tables, start=1):
            rows = table.rows[: self.limits.max_scan_rows]
            if not rows:
                continue
            raw_headers = tuple(
                _text_value(cell.text) for cell in rows[0].cells[: self.limits.max_scan_columns]
            )
            headers = _recognized_headers(raw_headers, variant, context_fields.keys())
            if headers is None:
                continue
            for row_index, row in enumerate(rows[1:], start=2):
                fields = tuple(
                    ObservedField(
                        name=header,
                        value=value,
                        source_ref=(
                            f"{checked_path.name}#table={table_index},row={row_index},"
                            f"column={column_index + 1}"
                        ),
                    )
                    for column_index, header in headers.by_column.items()
                    if column_index < len(row.cells)
                    if (value := _record_value(row.cells[column_index].text)) is not None
                )
                inherited_fields = tuple(
                    field
                    for canonical_field, fields_for_canonical in context_fields.items()
                    if canonical_field not in headers.canonical_fields
                    for field in fields_for_canonical
                )
                fields = inherited_fields + fields
                if fields:
                    records.append(
                        ObservedRecord(
                            record_ref=(f"{checked_path.name}#table={table_index},row={row_index}"),
                            fields=fields,
                        )
                    )
        return tuple(records)

    def extract_learning_structure(
        self, path: Path
    ) -> tuple[tuple[str, DocumentSignals], ...]:
        checked_path = _validate_office_archive(path, self.limits)
        try:
            document = Document(str(checked_path))
        except (
            OSError,
            ValueError,
            KeyError,
            SyntaxError,
            ParseError,
            OpcError,
            zipfile.BadZipFile,
        ) as error:
            raise OfficeDocumentError(f"Cannot parse DOCX document: {checked_path}") from error
        headings = tuple(
            value
            for paragraph in document.paragraphs[: self.limits.max_paragraphs]
            if (value := _text_value(paragraph.text)) is not None
            if (
                (paragraph.style.name if paragraph.style is not None else "")
                .casefold()
                .startswith("heading")
                or "見出し" in (paragraph.style.name if paragraph.style is not None else "")
            )
        )
        headers = tuple(
            value
            for table in document.tables
            if table.rows
            for value in _text_values(
                cell.text for cell in table.rows[0].cells[: self.limits.max_scan_columns]
            )
        )
        return (
            (
                "document",
                DocumentSignals.from_raw(
                    filename=checked_path.name,
                    headings=headings,
                    headers=headers,
                ),
            ),
        )


def _validate_office_archive(path: Path, limits: ExtractionLimits) -> Path:
    checked_path = path.resolve()
    if not checked_path.is_file():
        raise OfficeDocumentError(f"Document does not exist or is not a file: {path}")
    if checked_path.stat().st_size > limits.max_file_bytes:
        raise OfficeDocumentSecurityError(
            f"Office document exceeds {limits.max_file_bytes} bytes: {checked_path}"
        )

    try:
        with zipfile.ZipFile(checked_path) as archive:
            entries = archive.infolist()
            if len(entries) > limits.max_archive_entries:
                raise OfficeDocumentSecurityError(
                    f"Office archive has more than {limits.max_archive_entries} entries"
                )
            total_uncompressed = 0
            for entry in entries:
                _validate_archive_member(entry)
                total_uncompressed += entry.file_size
                if total_uncompressed > limits.max_uncompressed_bytes:
                    raise OfficeDocumentSecurityError(
                        "Office archive exceeds the configured uncompressed size limit"
                    )
    except zipfile.BadZipFile as error:
        raise OfficeDocumentSecurityError(
            f"Invalid Office ZIP container: {checked_path}"
        ) from error
    return checked_path


def _validate_archive_member(entry: zipfile.ZipInfo) -> None:
    normalized_name = entry.filename.replace("\\", "/")
    member_path = PurePosixPath(normalized_name)
    if (
        not normalized_name
        or "\x00" in normalized_name
        or member_path.is_absolute()
        or ".." in member_path.parts
    ):
        raise OfficeDocumentSecurityError(f"Unsafe Office archive member path: {entry.filename}")
    if entry.flag_bits & 0x1:
        raise OfficeDocumentSecurityError(f"Encrypted Office archive member: {entry.filename}")


def _text_values(values: Iterable[object]) -> list[str]:
    return [value for raw in values if (value := _text_value(raw)) is not None]


def _text_value(value: object) -> str | None:
    if not isinstance(value, str):
        return None
    stripped = value.strip()
    if not stripped or stripped.startswith("="):
        return None
    return stripped


def _record_value(value: object) -> str | None:
    text = _text_value(value)
    if text is not None:
        return text
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, datetime | date | time):
        return value.isoformat()
    if isinstance(value, timedelta):
        return str(value.total_seconds())
    if isinstance(value, int):
        return str(value)
    if isinstance(value, float) and isfinite(value):
        return str(int(value)) if value.is_integer() else repr(value)
    return None


def _recognized_headers(
    values: Iterable[str | None],
    variant: ConventionVariant,
    context_field_names: Collection[str],
) -> _DetectedHeaders | None:
    alias_lookup = _variant_alias_lookup(variant)

    headers: dict[int, str] = {}
    represented_fields: set[str] = set()
    for column_index, value in enumerate(values):
        if value is None:
            continue
        detected_canonical_field = alias_lookup.get(normalize_field_name(value))
        if detected_canonical_field is None:
            continue
        headers[column_index] = value
        represented_fields.add(detected_canonical_field)
    stable_fields = set(variant.stable_key_fields)
    if (
        len(headers) < 2
        or not represented_fields.intersection(stable_fields)
        or not stable_fields.issubset(represented_fields | set(context_field_names))
    ):
        return None
    return _DetectedHeaders(headers, frozenset(represented_fields))


def _variant_alias_lookup(variant: ConventionVariant) -> dict[str, str]:
    alias_lookup: dict[str, str] = {}
    for canonical_field, aliases in variant.field_aliases.items():
        for alias in aliases:
            normalized_alias = normalize_field_name(alias)
            existing = alias_lookup.get(normalized_alias)
            if existing is not None and existing != canonical_field:
                raise ValueError(f"Ambiguous Variant field alias: {alias}")
            alias_lookup[normalized_alias] = canonical_field
    return alias_lookup


def _key_value_context_fields(
    rows: Iterable[tuple[tuple[object, ...], tuple[str, ...]]],
    variant: ConventionVariant,
) -> dict[str, tuple[ObservedField, ...]]:
    return _key_value_fields(
        rows,
        variant,
        allowed_fields=frozenset(variant.stable_key_fields),
    )


def _key_value_fields(
    rows: Iterable[tuple[tuple[object, ...], tuple[str, ...]]],
    variant: ConventionVariant,
    *,
    allowed_fields: frozenset[str] | None = None,
) -> dict[str, tuple[ObservedField, ...]]:
    alias_lookup = _variant_alias_lookup(variant)
    collected: dict[str, list[ObservedField]] = {}
    materialized_rows = tuple(rows)
    for values, source_refs in materialized_rows:
        for column_index, raw_label in enumerate(values[:-1]):
            label = _text_value(raw_label)
            if label is None:
                continue
            canonical_field = alias_lookup.get(normalize_field_name(label))
            if canonical_field is None or (
                allowed_fields is not None and canonical_field not in allowed_fields
            ):
                continue
            value = _record_value(values[column_index + 1])
            if value is None:
                continue
            if normalize_field_name(value) in alias_lookup:
                continue
            collected.setdefault(canonical_field, []).append(
                ObservedField(
                    name=label,
                    value=value,
                    source_ref=source_refs[column_index + 1],
                )
            )
    for (values, _source_refs), (next_values, next_source_refs) in pairwise(materialized_rows):
        current = [
            (index, value)
            for index, raw in enumerate(values)
            if (value := _text_value(raw)) is not None
        ]
        following = [
            (index, value)
            for index, raw in enumerate(next_values)
            if (value := _record_value(raw)) is not None
        ]
        if len(current) != 1 or len(following) != 1:
            continue
        label = current[0][1]
        canonical_field = alias_lookup.get(normalize_field_name(label))
        if canonical_field is None or (
            allowed_fields is not None and canonical_field not in allowed_fields
        ):
            continue
        value_index, value = following[0]
        if normalize_field_name(value) in alias_lookup:
            continue
        collected.setdefault(canonical_field, []).append(
            ObservedField(
                name=label,
                value=value,
                source_ref=next_source_refs[value_index],
            )
        )
    if "operation" in variant.field_aliases and (
        allowed_fields is None or "operation" in allowed_fields
    ):
        operation_alias = variant.field_aliases["operation"][0]
        for values, source_refs in materialized_rows:
            for index, raw_value in enumerate(values):
                value = _text_value(raw_value)
                if value is None or HTTP_OPERATION.fullmatch(value) is None:
                    continue
                collected.setdefault("operation", []).append(
                    ObservedField(
                        name=operation_alias,
                        value=value,
                        source_ref=source_refs[index],
                    )
                )
    return {
        canonical_field: tuple(fields)
        for canonical_field, fields in collected.items()
        if len({normalize_business_value(field.value) for field in fields}) == 1
    }


def _section_marker_fields(
    values: tuple[str | None, ...],
    source_refs: tuple[str, ...],
    variant: ConventionVariant,
    *,
    excluded_fields: frozenset[str],
) -> dict[str, ObservedField]:
    non_empty = [(index, value) for index, value in enumerate(values) if value is not None]
    if len(non_empty) != 1:
        return {}
    index, value = non_empty[0]
    canonical_field = _variant_alias_lookup(variant).get(normalize_field_name(value))
    if canonical_field is None or canonical_field in excluded_fields:
        return {}
    if canonical_field not in variant.stable_key_fields:
        return {}
    return {
        canonical_field: ObservedField(
            name=value,
            value=value,
            source_ref=source_refs[index],
        )
    }
