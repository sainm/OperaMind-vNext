from pathlib import Path
from types import SimpleNamespace
from typing import Any, cast

from docx import Document
from openpyxl import Workbook

from operamind.application import project_document_baseline as baseline_module
from operamind.application.project_document_baseline import ProjectDocumentBaselineService

ROOT = Path(__file__).parents[2]


def _service() -> ProjectDocumentBaselineService:
    return ProjectDocumentBaselineService(
        connection=cast(Any, object()),
        repository_root=ROOT,
    )


def test_discovery_uses_profile_signals_instead_of_japanese_filename(tmp_path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Screen Items"
    sheet.append(["Screen ID", "Element ID", "Type", "Default Value", "Notes"])
    sheet.append(["customer-list", "status-filter", "select", "all", "filter"])
    path = tmp_path / "customer-ui-specification.xlsx"
    workbook.save(path)

    discovery = _service().discover(document_roots=(tmp_path,))

    assert discovery.ready is True
    assert len(discovery.candidates) == 1
    candidate = discovery.candidates[0]
    assert candidate.path == path.resolve()
    assert candidate.profile_id == "screen-design-conventions-example"
    assert candidate.document_type == "screen_design"
    assert candidate.fact_type == "screen_element"
    assert candidate.score == 0.8


def test_discovery_supports_docx_through_the_same_profile_registry(tmp_path: Path) -> None:
    document = Document()
    document.add_heading("Interface Specification", level=1)
    table = document.add_table(rows=2, cols=3)
    for cell, value in zip(table.rows[0].cells, ("Path", "Operation", "Response"), strict=True):
        cell.text = value
    for cell, value in zip(table.rows[1].cells, ("/expenses", "GET", "ExpenseList"), strict=True):
        cell.text = value
    path = tmp_path / "external-contract.docx"
    document.save(path)

    discovery = _service().discover(document_roots=(tmp_path,))

    assert discovery.ready is True
    assert len(discovery.candidates) == 1
    candidate = discovery.candidates[0]
    assert candidate.path == path.resolve()
    assert candidate.profile_id == "api-design-conventions-example"
    assert candidate.document_type == "api_design"
    assert candidate.fact_type == "api_endpoint"
    assert candidate.score == 1.0


def test_discovery_reports_partial_matches_for_review_and_ignores_unrelated_docs(
    tmp_path: Path,
) -> None:
    partial = Document()
    partial.add_heading("Interface Specification", level=1)
    partial.save(tmp_path / "partial.docx")
    unrelated = Document()
    unrelated.add_paragraph("Meeting notes without a design table")
    unrelated.save(tmp_path / "notes.docx")

    discovery = _service().discover(document_roots=(tmp_path,))

    assert discovery.ready is False
    assert discovery.candidates == ()
    assert discovery.ignored_documents == (str((tmp_path / "notes.docx").resolve()),)
    assert len(discovery.review_required) == 1
    assert "partial.docx" in discovery.review_required[0]
    assert "api-design-conventions-example" in discovery.review_required[0]


def test_discovery_summary_is_business_safe_and_contains_no_profile_payload(tmp_path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Method List"
    sheet.append(["Method Name", "Arguments", "Return Type", "Summary"])
    sheet.append(["search", "status", "List", "Search expenses"])
    workbook.save(tmp_path / "service-design.xlsx")

    summary = _service().discover(document_roots=(tmp_path,)).public_summary()

    assert summary["status"] == "ready"
    assert summary["document_count"] == 1
    assert summary["documents"] == [
        {
            "path": str((tmp_path / "service-design.xlsx").resolve()),
            "profile_id": "program-design-conventions-example",
            "document_type": "program_design",
            "fact_type": "program_method",
            "match_score": 0.8,
        }
    ]
    assert "variants" not in repr(summary)


def test_store_documents_persists_profile_backed_snapshot_and_nodes(tmp_path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Screen Items"
    sheet.append(["Screen ID", "Element ID", "Type", "Default Value", "Notes"])
    sheet.append(["customer-list", "status-filter", "select", "all", "filter"])
    workbook.save(tmp_path / "customer-ui-specification.xlsx")

    service = _service()
    snapshots: list[object] = []
    stored_nodes: list[tuple[object, ...]] = []
    activations: list[dict[str, object]] = []

    class ProfileRepository:
        def store_version(self, **_values: object) -> str:
            return "profile-digest"

        def activate(self, **values: object) -> None:
            activations.append(values)

    service._profile_repository = cast(Any, ProfileRepository())
    service._canonical = cast(
        Any,
        SimpleNamespace(store_snapshot=lambda snapshot: snapshots.append(snapshot)),
    )
    service._nodes = cast(
        Any,
        SimpleNamespace(
            store_nodes=lambda **values: stored_nodes.append(tuple(values["nodes"]))
        ),
    )
    discovery = service.discover(document_roots=(tmp_path,))
    service.discover = cast(Any, lambda **_values: discovery)

    result = service.store_documents(
        project_id="profile-project",
        document_roots=(tmp_path,),
        actor="operator",
    )

    assert result.document_count == 1
    assert result.snapshot_id.startswith("document-baseline-")
    assert len(snapshots) == 1
    assert len(stored_nodes) == 1 and stored_nodes[0]
    assert activations[0]["binding_key"] == "document:screen_design"


def test_build_index_uses_snapshot_identity_and_returns_published_result(
    monkeypatch: Any,
) -> None:
    captured: dict[str, object] = {}

    class IndexService:
        def __init__(self, **values: object) -> None:
            captured["constructor"] = values

        def run(self, request: object, **values: object) -> SimpleNamespace:
            captured["request"] = request
            captured["run"] = values
            return SimpleNamespace(
                state=SimpleNamespace(spec=SimpleNamespace(build_id=request.build_id)),
                generated_vector_count=3,
            )

    provider = object()
    monkeypatch.setattr(baseline_module, "SearchIndexBuildService", IndexService)
    monkeypatch.setattr(
        baseline_module.OpenAICompatibleEmbeddingProvider,
        "from_profile",
        lambda _profile: provider,
    )

    result = _service().build_index(
        project_id="profile-project",
        snapshot_id="snapshot-current",
        document_count=2,
        actor="operator",
        build_nonce="attempt-2",
    )

    assert result.snapshot_id == "snapshot-current"
    assert result.document_count == 2
    assert result.generated_vector_count == 3
    assert result.index_build_id.startswith("search-index-")
    assert captured["run"]["provider"] is provider
