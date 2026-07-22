import json
import zipfile
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from openpyxl import Workbook

from operamind.domain import CanonicalFactMapper, CanonicalMappingStatus
from operamind.domain.document_conventions import (
    ConventionMatcher,
    DocumentConvention,
    MatchStatus,
)
from operamind.infrastructure.documents import (
    DocumentSignalExtractorRegistry,
    ExtractionLimits,
    OfficeDocumentError,
    OfficeDocumentSecurityError,
    UnsupportedDocumentTypeError,
)
from operamind.profiles import ProfileCatalog

ROOT = Path(__file__).parents[3]


def load_convention() -> DocumentConvention:
    profile: dict[str, Any] = json.loads(
        (ROOT / "profiles/document-convention-profile.example.json").read_text(encoding="utf-8")
    )
    ProfileCatalog.load(ROOT / "profiles").validate_profile(profile)
    return DocumentConvention.from_validated_profile(profile)


def test_xlsx_extraction_feeds_variant_matcher(tmp_path: Path) -> None:
    path = tmp_path / "顧客_API_設計書.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "API一覧"
    worksheet.append(["API設計"])
    worksheet.append(["URI", "HTTPメソッド", "処理概要"])
    worksheet.append(["/expenses", "GET", "経費一覧"])
    workbook.save(path)
    workbook.close()

    signals = DocumentSignalExtractorRegistry.default().extract(path)
    extractor_ref = DocumentSignalExtractorRegistry.default().extractor_ref(path)
    result = ConventionMatcher().match(load_convention(), signals)
    variant = next(
        variant
        for variant in load_convention().variants
        if variant.variant_id == result.selected_variant_id
    )
    records = DocumentSignalExtractorRegistry.default().extract_records(path, variant)
    mapping = CanonicalFactMapper().map_record(
        convention=load_convention(),
        match=result,
        fact_type="api",
        record=records[0],
    )

    assert "api一覧" in signals.sheet_names
    assert "api設計" in signals.headings
    assert {"uri", "httpメソッド", "処理概要"} <= signals.headers
    assert "経費一覧" in signals.business_terms
    assert extractor_ref.startswith("operamind-xlsx-structural@1+openpyxl-")
    assert result.status is MatchStatus.AUTO_MATCHED
    assert result.selected_variant_id == "api-list"
    assert len(records) == 1
    assert mapping.status is CanonicalMappingStatus.MAPPED
    assert mapping.fact is not None
    assert mapping.fact.stable_key == "api:GET/%2Fexpenses"
    assert mapping.fact.source_refs == (
        "顧客_API_設計書.xlsx#API一覧!A3",
        "顧客_API_設計書.xlsx#API一覧!B3",
        "顧客_API_設計書.xlsx#API一覧!C3",
    )


def test_docx_extraction_reads_headings_and_first_table_row(tmp_path: Path) -> None:
    path = tmp_path / "interface.docx"
    document = Document()
    document.add_heading("インターフェース仕様", level=1)
    document.add_paragraph("顧客情報を取得する API")
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Path"
    table.rows[0].cells[1].text = "Operation"
    table.rows[0].cells[2].text = "Response"
    table.rows[1].cells[0].text = "/customers"
    table.rows[1].cells[1].text = "GET"
    table.rows[1].cells[2].text = "Customer"
    document.save(path)

    signals = DocumentSignalExtractorRegistry.default().extract(path)
    extractor_ref = DocumentSignalExtractorRegistry.default().extractor_ref(path)
    convention = load_convention()
    match = ConventionMatcher().match(convention, signals)
    variant = next(
        variant
        for variant in convention.variants
        if variant.variant_id == match.selected_variant_id
    )
    records = DocumentSignalExtractorRegistry.default().extract_records(path, variant)
    mapping = CanonicalFactMapper().map_record(
        convention=convention,
        match=match,
        fact_type="api",
        record=records[0],
    )

    assert "インターフェース仕様" in signals.headings
    assert {"path", "operation", "response"} <= signals.headers
    assert "顧客情報を取得する api" in signals.business_terms
    assert "customer" in signals.business_terms
    assert extractor_ref.startswith("operamind-docx-structural@1+python-docx-")
    assert len(records) == 1
    assert mapping.status is CanonicalMappingStatus.MAPPED
    assert mapping.fact is not None
    assert mapping.fact.stable_key == "api:GET/%2Fcustomers"


def test_xlsx_layout_and_alias_changes_preserve_canonical_fact(tmp_path: Path) -> None:
    before_path = tmp_path / "API-before.xlsx"
    after_path = tmp_path / "API-after.xlsx"
    before = Workbook()
    before.active.title = "API一覧"
    before.active.append(["URI", "HTTPメソッド", "処理概要"])
    before.active.append(["/expenses", "GET", "経費一覧"])
    before.save(before_path)
    before.close()
    after = Workbook()
    after.active.title = "Renamed Sheet"
    after.active.append(["Description", "Method", "Path"])
    after.active.append(["経費一覧", "GET", "/expenses"])
    after.save(after_path)
    after.close()

    convention = load_convention()
    match = ConventionMatcher().match(
        convention,
        DocumentSignalExtractorRegistry.default().extract(before_path),
    )
    variant = next(
        variant
        for variant in convention.variants
        if variant.variant_id == match.selected_variant_id
    )
    registry = DocumentSignalExtractorRegistry.default()
    before_record = registry.extract_records(before_path, variant)[0]
    after_record = registry.extract_records(after_path, variant)[0]
    mapper = CanonicalFactMapper()

    before_result = mapper.map_record(
        convention=convention,
        match=match,
        fact_type="api",
        record=before_record,
    )
    after_result = mapper.map_record(
        convention=convention,
        match=match,
        fact_type="api",
        record=after_record,
    )

    assert before_result.fact is not None
    assert after_result.fact is not None
    assert before_result.fact.stable_key == after_result.fact.stable_key
    assert before_result.fact.values == after_result.fact.values


@pytest.mark.parametrize("suffix", [".xls", ".doc", ".xlsm", ""])
def test_registry_rejects_unregistered_document_types(tmp_path: Path, suffix: str) -> None:
    path = tmp_path / f"legacy{suffix}"
    path.write_bytes(b"not parsed")

    with pytest.raises(UnsupportedDocumentTypeError, match="Unsupported document type"):
        DocumentSignalExtractorRegistry.default().extract(path)


def test_archive_member_path_traversal_is_rejected_before_parsing(tmp_path: Path) -> None:
    path = tmp_path / "unsafe.xlsx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("../outside.xml", "unsafe")

    with pytest.raises(OfficeDocumentSecurityError, match="Unsafe Office archive member path"):
        DocumentSignalExtractorRegistry.default().extract(path)


def test_uncompressed_archive_limit_is_enforced(tmp_path: Path) -> None:
    path = tmp_path / "oversized.xlsx"
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("xl/workbook.xml", "x" * 128)
    limits = ExtractionLimits(max_uncompressed_bytes=64)

    with pytest.raises(OfficeDocumentSecurityError, match="uncompressed size limit"):
        DocumentSignalExtractorRegistry.default(limits).extract(path)


@pytest.mark.parametrize("suffix", [".xlsx", ".docx"])
def test_malformed_ooxml_is_wrapped_as_document_error(tmp_path: Path, suffix: str) -> None:
    path = tmp_path / f"malformed{suffix}"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("placeholder.xml", "<not-office />")

    with pytest.raises(OfficeDocumentError, match="Cannot parse"):
        DocumentSignalExtractorRegistry.default().extract(path)


def test_extraction_limits_must_be_positive() -> None:
    with pytest.raises(ValueError, match="max_scan_rows"):
        ExtractionLimits(max_scan_rows=0)
